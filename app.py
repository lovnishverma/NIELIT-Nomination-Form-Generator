import logging
import os
import re
import tempfile
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document
from flask import Flask, flash, redirect, render_template, request, send_file

# ─── Logging ────────────────────────────────────────────────────────────────
logging.basicConfig(
    filename="app_errors.log",
    level=logging.ERROR,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

BASE_DIR = Path(__file__).resolve().parent
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev_secret_key_change_me")
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB

# ─── 2-Template Architecture ─────────────────────────────────────────────────
# One GOT template covers BASIC + ADVANCED for both ARVR and BDDS.
# One BOOTCAMP template covers Bootcamp for both ARVR and BDDS.
# Dynamic fields (Technology, Domain, Course_Name …) distinguish them at runtime.
TEMPLATE_MAP = {
    "GOT":      BASE_DIR / "docxtemplates" / "GOT_Nomination_Form.docx",
    "BOOTCAMP": BASE_DIR / "docxtemplates" / "Bootcamp_Nomination_Form.docx",
}

# Which template key each logical form type uses
FORM_TYPE_TEMPLATE = {
    "ARVR_BASIC":    "GOT",
    "ARVR_ADVANCED": "GOT",
    "BDDS_BASIC":    "GOT",
    "BDDS_ADVANCED": "GOT",
    "ARVR_BOOTCAMP": "BOOTCAMP",
    "BDDS_BOOTCAMP": "BOOTCAMP",
}

# Canonical aliases  →  logical form type
FORM_TYPE_ALIASES = {
    "ARVR_BASIC":         "ARVR_BASIC",
    "ARVR_GOT_BASIC":     "ARVR_BASIC",
    "ARVR_ADVANCED":      "ARVR_ADVANCED",
    "ARVR_GOT_ADVANCED":  "ARVR_ADVANCED",
    "ARVR_BOOTCAMP":      "ARVR_BOOTCAMP",
    "ARVR_BOOT_CAMP":     "ARVR_BOOTCAMP",
    "ARVR_GOT_BOOTCAMP":  "ARVR_BOOTCAMP",
    "BDDS_BASIC":         "BDDS_BASIC",
    "BDDS_GOT_BASIC":     "BDDS_BASIC",
    "BDDS_ADVANCED":      "BDDS_ADVANCED",
    "BDDS_GOT_ADVANCED":  "BDDS_ADVANCED",
    "BDDS_BOOTCAMP":      "BDDS_BOOTCAMP",
    "BDDS_BOOT_CAMP":     "BDDS_BOOTCAMP",
    "BDDS_GOT_BOOTCAMP":  "BDDS_BOOTCAMP",
}

# Per-form-type defaults injected if the CSV row leaves them blank
FORM_TYPE_DEFAULTS = {
    "ARVR_BASIC": {
        "Domain":           "Augmented and Virtual Reality (AR & VR)",
        "Course_Name":      "GOT – AR & VR (Basic)",
        "Course_Level":     "Basic",
        "Technology":       "Augmented and Virtual Reality",
        "Program_Type":     "GOT (Government of Tomorrow)",
        "Form_Title":       "ARVR GOT-Basic Nomination Form",
    },
    "ARVR_ADVANCED": {
        "Domain":           "Augmented and Virtual Reality (AR & VR)",
        "Course_Name":      "GOT – AR & VR (Advanced)",
        "Course_Level":     "Advanced",
        "Technology":       "Augmented and Virtual Reality",
        "Program_Type":     "GOT (Government of Tomorrow)",
        "Form_Title":       "ARVR GOT-Advanced Nomination Form",
    },
    "ARVR_BOOTCAMP": {
        "Domain":           "Augmented and Virtual Reality (AR & VR)",
        "Course_Name":      "Bootcamp – AR & VR",
        "Course_Level":     "Bootcamp",
        "Technology":       "Augmented and Virtual Reality",
        "Program_Type":     "Bootcamp",
        "Form_Title":       "ARVR Bootcamp Nomination Form",
    },
    "BDDS_BASIC": {
        "Domain":           "Big Data & Data Science",
        "Course_Name":      "GOT – Big Data & Data Science (Basic)",
        "Course_Level":     "Basic",
        "Technology":       "Big Data & Data Science",
        "Program_Type":     "GOT (Government of Tomorrow)",
        "Form_Title":       "BDDS GOT-Basic Nomination Form",
    },
    "BDDS_ADVANCED": {
        "Domain":           "Big Data & Data Science",
        "Course_Name":      "GOT – Big Data & Data Science (Advanced)",
        "Course_Level":     "Advanced",
        "Technology":       "Big Data & Data Science",
        "Program_Type":     "GOT (Government of Tomorrow)",
        "Form_Title":       "BDDS GOT-Advanced Nomination Form",
    },
    "BDDS_BOOTCAMP": {
        "Domain":           "Big Data & Data Science",
        "Course_Name":      "Bootcamp – Big Data & Data Science",
        "Course_Level":     "Bootcamp",
        "Technology":       "Big Data & Data Science",
        "Program_Type":     "Bootcamp",
        "Form_Title":       "BDDS Bootcamp Nomination Form",
    },
}

COLUMN_ALIASES = {
    "AADHAR_NO":          "Aadhar",
    "AADHAAR_NO":         "Aadhar",
    "AADHAAR_NUMBER":     "Aadhar",
    "MOBILE_NUMBER":      "Contact_Number",
    "PHONE_NUMBER":       "Contact_Number",
    "EMAIL_ID":           "Email",
    "E_MAIL":             "Email",
    "NAME_OF_THE_ORGANISATION_DEPARTMENT": "Organisation_Department",
    "COMPLETE_ADDRESS_CONTACT_NUMBERS_E_MAIL_OF_THE_INSTITUTE":
        "Institute_Address_Contact_Email",
}


# ─── Helpers ─────────────────────────────────────────────────────────────────

def normalize_token(value: str) -> str:
    return re.sub(r"[^A-Z0-9]+", "_", str(value).upper()).strip("_")


def infer_form_type(row_data: dict):
    """
    Priority order:
      1. Form_Type column
      2. Track + Level columns
      3. Domain/Program + Level variants
    Returns a canonical key like 'ARVR_BASIC' or None.
    """
    raw = str(row_data.get("Form_Type", "")).strip()
    if raw:
        key = FORM_TYPE_ALIASES.get(normalize_token(raw))
        if key:
            return key

    track = (
        row_data.get("Track") or row_data.get("Domain") or
        row_data.get("Program") or row_data.get("Course") or
        row_data.get("Course_Track") or ""
    )
    level = (
        row_data.get("Level") or row_data.get("Course_Level") or
        row_data.get("Batch_Type") or row_data.get("Training_Level") or ""
    )
    if track and level:
        combined = f"{normalize_token(track)}_{normalize_token(level)}"
        key = FORM_TYPE_ALIASES.get(combined)
        if key:
            return key

    return None


def sanitize_filename(value, fallback="Applicant"):
    cleaned = re.sub(r"[^a-zA-Z0-9_\-]+", "_", str(value).strip()).strip("_")
    return cleaned or fallback


def unique_output_name(base_name, form_type, used_names):
    candidate = f"{base_name}_{form_type}.docx"
    if candidate not in used_names:
        used_names.add(candidate)
        return candidate
    counter = 2
    while True:
        candidate = f"{base_name}_{form_type}_{counter}.docx"
        if candidate not in used_names:
            used_names.add(candidate)
            return candidate
        counter += 1


def _safe(d, key):
    return str(d.get(key, "")).strip()


def _merge_non_empty(parts, sep=" | "):
    return sep.join(p for p in parts if p)


def _build_education_block(row_data):
    lines = []
    for i in range(1, 6):
        year   = _safe(row_data, f"Edu{i}_Year")
        degree = _safe(row_data, f"Edu{i}_Degree")
        univ   = _safe(row_data, f"Edu{i}_University")
        if any([year, degree, univ]):
            lines.append(f"{i}. {year} – {degree} – {univ}".strip(" –"))
    return "\n".join(lines)


def _build_experience_block(row_data):
    lines = []
    for i in range(1, 6):
        year   = _safe(row_data, f"Exp{i}_Year")
        area   = _safe(row_data, f"Exp{i}_Area_of_Expertise")
        centre = _safe(row_data, f"Exp{i}_Centre")
        if any([year, area, centre]):
            lines.append(f"{i}. {year} – {area} – {centre}".strip(" –"))
    return "\n".join(lines)


def enrich_row_data(row_data: dict, form_type: str) -> dict:
    """
    1. Normalise column aliases.
    2. Inject per-form-type defaults for domain/course fields.
    3. Auto-build composite fields when missing.
    """
    # Step 1 – alias normalisation
    normalized = {}
    for key, value in row_data.items():
        key_str = str(key).strip()
        normalized[key_str] = value
        alias = COLUMN_ALIASES.get(normalize_token(key_str))
        if alias and not normalized.get(alias):
            normalized[alias] = value

    # Step 2 – form-type defaults (only fill blanks, never overwrite user data)
    for field, default_val in FORM_TYPE_DEFAULTS.get(form_type, {}).items():
        if not normalized.get(field):
            normalized[field] = default_val

    # Step 3 – composite / derived fields
    normalized.setdefault("Organisation_Department",
        _merge_non_empty([_safe(normalized, "Organisation"),
                          _safe(normalized, "Department")], " / "))

    normalized.setdefault("Contact_Number_Email",
        _merge_non_empty([_safe(normalized, "Contact_Number"),
                          _safe(normalized, "Email")], " / "))

    normalized.setdefault("Institute_Address_Contact_Email",
        _merge_non_empty([_safe(normalized, "Institute_Address"),
                          _safe(normalized, "Institute_Contact"),
                          _safe(normalized, "Institute_Email")], " / "))

    if not normalized.get("Educational_Qualifications"):
        normalized["Educational_Qualifications"] = _build_education_block(normalized)

    if not normalized.get("Research_Technical_Experience"):
        normalized["Research_Technical_Experience"] = _build_experience_block(normalized)

    normalized.setdefault("Applicant_Name",
        normalized.get("Name") or normalized.get("Candidate_Name") or "")

    normalized.setdefault("Gov_ID_Number",
        normalized.get("Aadhar") or "")

    # Friendly date today for signing fields
    normalized.setdefault("Today_Date", datetime.today().strftime("%d-%m-%Y"))

    return normalized


# ─── DOCX replacement ────────────────────────────────────────────────────────

def _build_replacements(data: dict) -> dict:
    """Build {key: value} for every placeholder variant."""
    replacements = {}
    for key, value in data.items():
        key_str   = str(key).strip()
        value_str = "" if value is None else str(value)
        for variant in {key_str, key_str.upper(), key_str.lower()}:
            replacements[variant] = value_str
    return replacements


def _replace_in_paragraph(paragraph, replacements: dict):
    """Replace placeholders run-by-run to preserve formatting."""
    full_text = paragraph.text
    # Quick bail-out – nothing to replace
    if "{" not in full_text and "<<" not in full_text:
        return

    if paragraph.runs:
        # First pass: try to replace within individual runs
        for run in paragraph.runs:
            if not run.text:
                continue
            updated = run.text
            for ph, val in replacements.items():
                updated = updated.replace(f"{{{ph}}}", val)
                updated = updated.replace(f"<<{ph}>>", val)
            if updated != run.text:
                run.text = updated
        # Second pass: check if a placeholder was split across runs
        # If so, collapse and re-expand
        joined = "".join(r.text for r in paragraph.runs)
        still_has_ph = "{" in joined or "<<" in joined
        if still_has_ph:
            # Replace in joined text
            updated_joined = joined
            for ph, val in replacements.items():
                updated_joined = updated_joined.replace(f"{{{ph}}}", val)
                updated_joined = updated_joined.replace(f"<<{ph}>>", val)
            if updated_joined != joined:
                # Write everything into first run, clear the rest
                paragraph.runs[0].text = updated_joined
                for run in paragraph.runs[1:]:
                    run.text = ""
    else:
        # No runs – paragraph.text setter fallback
        updated = full_text
        for ph, val in replacements.items():
            updated = updated.replace(f"{{{ph}}}", val)
            updated = updated.replace(f"<<{ph}>>", val)
        if updated != full_text:
            paragraph.text = updated


def fill_form(template_path: Path, output_path: Path, data: dict) -> bool:
    """Fill all placeholders in the template and save to output_path."""
    try:
        doc = Document(str(template_path))
        replacements = _build_replacements(data)

        for para in doc.paragraphs:
            _replace_in_paragraph(para, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, replacements)

        # Also handle headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                _replace_in_paragraph(para, replacements)
            for para in section.footer.paragraphs:
                _replace_in_paragraph(para, replacements)

        doc.save(str(output_path))
        return True
    except Exception as exc:
        logging.error(
            f"fill_form failed – template={template_path.name}, "
            f"name={data.get('Name', 'Unknown')}: {exc}"
        )
        return False


# ─── CSV reading ─────────────────────────────────────────────────────────────

def read_csv_safely(file_storage) -> pd.DataFrame:
    raw = file_storage.read()
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return pd.read_csv(BytesIO(raw), encoding=enc)
        except Exception:
            continue
    raise ValueError("Could not decode CSV. Please save as UTF-8 and try again.")


# ─── Routes ──────────────────────────────────────────────────────────────────

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files.get("file")

        if not file or file.filename == "":
            flash("Please select a CSV file.", "error")
            return redirect(request.url)

        if not file.filename.lower().endswith(".csv"):
            flash("Only CSV files are supported.", "error")
            return redirect(request.url)

        try:
            try:
                df = read_csv_safely(file)
            except Exception as exc:
                logging.error(f"CSV read error: {exc}")
                flash("Failed to read the CSV file. Ensure it is not corrupted.", "error")
                return redirect(request.url)

            if df.empty:
                flash("The CSV file is empty.", "error")
                return redirect(request.url)

            df.columns = [str(c).strip() for c in df.columns]

            timestamp     = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            memory_zip    = BytesIO()
            success_count = 0
            error_count   = 0
            used_names    = set()
            skipped_rows  = []

            with tempfile.TemporaryDirectory() as tmp:
                with zipfile.ZipFile(memory_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                    for idx, row in df.iterrows():
                        row_data  = row.fillna("").to_dict()
                        form_type = infer_form_type(row_data)

                        if not form_type:
                            msg = (
                                f"Row {idx + 2}: Cannot determine form type. "
                                "Add a Form_Type column (e.g. ARVR_BASIC) "
                                "or both Track and Level columns."
                            )
                            logging.error(msg)
                            skipped_rows.append(f"Row {idx + 2}: unknown form type")
                            error_count += 1
                            continue

                        template_key  = FORM_TYPE_TEMPLATE[form_type]
                        template_path = TEMPLATE_MAP[template_key]

                        if not template_path.exists():
                            msg = (
                                f"Row {idx + 2}: Template file missing – "
                                f"{template_path.name}. "
                                "Place it in the docxtemplates/ folder."
                            )
                            logging.error(msg)
                            skipped_rows.append(f"Row {idx + 2}: template missing ({template_key})")
                            error_count += 1
                            continue

                        enriched = enrich_row_data(row_data, form_type)
                        name     = (
                            enriched.get("Name") or
                            enriched.get("Candidate_Name") or
                            f"Applicant_{idx + 2}"
                        )
                        safe_name    = sanitize_filename(name, f"Applicant_{idx + 2}")
                        out_filename = unique_output_name(safe_name, form_type, used_names)
                        out_path     = Path(tmp) / out_filename

                        if fill_form(template_path, out_path, enriched):
                            zf.write(str(out_path), arcname=out_filename)
                            success_count += 1
                        else:
                            skipped_rows.append(f"Row {idx + 2} ({name}): fill_form error")
                            error_count += 1

            memory_zip.seek(0)

            if success_count == 0:
                flash(
                    f"No forms generated. {error_count} row(s) failed. "
                    "Check app_errors.log for details.",
                    "error",
                )
                return redirect(request.url)

            if error_count > 0:
                flash(
                    f"Partial success: {success_count} form(s) generated, "
                    f"{error_count} row(s) skipped. See app_errors.log.",
                    "warning",
                )

            # Read the cookie token the client sent so we can echo it back.
            # The JS polls for this cookie to know the download has started
            # and hides the loading overlay.
            download_token = request.form.get("downloadToken", "done")
            response = send_file(
                memory_zip,
                mimetype="application/zip",
                as_attachment=True,
                download_name=f"Nomination_Batch_{timestamp}.zip",
            )
            response.set_cookie(
                "fileDownload", download_token,
                max_age=60, samesite="Lax"
            )
            return response

        except Exception as exc:
            logging.error(f"Unexpected error: {exc}")
            flash("A critical error occurred. Check app_errors.log.", "error")
            return redirect(request.url)

    template_status = {
        "GOT":      TEMPLATE_MAP["GOT"].exists(),
        "BOOTCAMP": TEMPLATE_MAP["BOOTCAMP"].exists(),
    }
    return render_template("index.html", template_status=template_status)


@app.errorhandler(413)
def file_too_large(_):
    flash("File too large. Maximum upload size is 10 MB.", "error")
    return redirect("/")


if __name__ == "__main__":
    app.run(debug=True)
