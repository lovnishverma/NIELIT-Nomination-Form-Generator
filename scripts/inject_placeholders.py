"""
inject_placeholders.py
──────────────────────
Standalone utility to test placeholder injection into a DOCX template
without running the Flask app.

Usage:
    python scripts/inject_placeholders.py \
        --template docxtemplates/GOT_Nomination_Form.docx \
        --csv sample_master.csv \
        --row 0 \
        --out test_output.docx

    # Or pass key=value pairs directly:
    python scripts/inject_placeholders.py \
        --template docxtemplates/Bootcamp_Nomination_Form.docx \
        --data "Name=Amit Kumar" "Form_Type=ARVR_BOOTCAMP" "Organisation=NIELIT"
"""

import argparse
import re
import sys
from pathlib import Path

# Allow running from project root
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pandas as pd
from docx import Document

from app import enrich_row_data, fill_form, infer_form_type, FORM_TYPE_TEMPLATE, TEMPLATE_MAP


def list_placeholders(doc_path: Path):
    """Print all {placeholder} and <<placeholder>> tokens found in a DOCX."""
    doc = Document(str(doc_path))
    found = set()
    pattern = re.compile(r"\{([^}]+)\}|<<([^>]+)>>")

    def _scan(text):
        for m in pattern.finditer(text):
            found.add(m.group(1) or m.group(2))

    for para in doc.paragraphs:
        _scan(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan(para.text)

    if found:
        print(f"\nPlaceholders found in {doc_path.name}:")
        for ph in sorted(found):
            print(f"  {{{ph}}}")
    else:
        print(f"No placeholders found in {doc_path.name}.")
    return found


def main():
    parser = argparse.ArgumentParser(description="Inject placeholders into a DOCX template.")
    parser.add_argument("--template", help="Path to .docx template (overrides auto-detection)")
    parser.add_argument("--csv", help="Path to CSV file")
    parser.add_argument("--row", type=int, default=0, help="Row index (0-based) to use from CSV")
    parser.add_argument("--out", default="test_output.docx", help="Output .docx path")
    parser.add_argument("--list", action="store_true", help="List placeholders in template and exit")
    parser.add_argument("--data", nargs="*", metavar="KEY=VALUE",
                        help="Manual key=value pairs (used if --csv not provided)")
    args = parser.parse_args()

    # ── Build row_data ──────────────────────────────────────────────────────
    if args.csv:
        df = pd.read_csv(args.csv, encoding="utf-8-sig")
        if args.row >= len(df):
            print(f"Error: CSV only has {len(df)} rows (0-based index {args.row} out of range).")
            sys.exit(1)
        row_data = df.iloc[args.row].fillna("").to_dict()
        print(f"Using row {args.row}: {row_data.get('Name', row_data.get('Candidate_Name', '?'))}")
    elif args.data:
        row_data = {}
        for item in args.data:
            if "=" in item:
                k, v = item.split("=", 1)
                row_data[k.strip()] = v.strip()
    else:
        print("Provide --csv or --data arguments.")
        parser.print_help()
        sys.exit(1)

    # ── Detect form type ────────────────────────────────────────────────────
    form_type = infer_form_type(row_data)
    if not form_type:
        print("Could not infer form type. Ensure Form_Type or Track+Level columns exist.")
        sys.exit(1)
    print(f"Detected form type: {form_type}")

    # ── Resolve template ────────────────────────────────────────────────────
    if args.template:
        template_path = Path(args.template)
    else:
        template_key  = FORM_TYPE_TEMPLATE[form_type]
        template_path = TEMPLATE_MAP[template_key]
        print(f"Auto-selected template: {template_path.name} ({template_key})")

    if not template_path.exists():
        print(f"Template not found: {template_path}")
        sys.exit(1)

    # ── List mode ───────────────────────────────────────────────────────────
    if args.list:
        list_placeholders(template_path)
        sys.exit(0)

    # ── Enrich + fill ───────────────────────────────────────────────────────
    enriched = enrich_row_data(row_data, form_type)
    out_path = Path(args.out)
    success  = fill_form(template_path, out_path, enriched)

    if success:
        print(f"\n✓ Output written to: {out_path}")
        print("\nFields injected:")
        for k, v in sorted(enriched.items()):
            if v:
                print(f"  {k}: {v}")
    else:
        print("✗ fill_form failed. Check app_errors.log.")
        sys.exit(1)


if __name__ == "__main__":
    main()
